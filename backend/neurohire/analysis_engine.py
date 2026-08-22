import fitz
import re
import os
import logging
import requests
import time
import json
from typing import Dict, List, Optional, Tuple
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import threading

# Setup logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize Sentence Transformer model globally
semantic_model = None
semantic_model_lock = threading.Lock()


def _get_semantic_model():
    """Load the Sentence Transformer model lazily on first use."""
    global semantic_model
    if semantic_model is not None:
        return semantic_model

    with semantic_model_lock:
        if semantic_model is not None:
            return semantic_model

        try:
            from sentence_transformers import SentenceTransformer

            semantic_model = SentenceTransformer('all-MiniLM-L6-v2')
            return semantic_model
        except Exception as e:
            logger.error(f"Failed to load Sentence Transformer model: {e}")
            semantic_model = None
            return None


# ==============================
# Gemini Structured Extraction
# ==============================

# Lazy, thread-safe client init (mirrors the semantic model pattern above)
gemini_client = None
gemini_client_lock = threading.Lock()

# gemini-2.5-flash and gemini-2.0-flash were both deprecated by Google after
# this code was originally written (2.0-flash shut down entirely in June 2026;
# 2.5-flash was restricted from new users ahead of its October 2026 shutdown).
# Updated to the current stable Flash generation as of August 2026.
GEMINI_MODEL_PRIMARY = 'gemini-3.6-flash'
GEMINI_MODEL_FALLBACK = 'gemini-3.5-flash'
GEMINI_TIMEOUT_SECONDS = 15  # keep tight — this runs inline in the upload request


def _get_gemini_client():
    """
    Load the Gemini client lazily on first use.
    Returns None (never raises) if the SDK is missing or no API key is configured —
    callers must treat None as "fall back to regex extraction".
    """
    global gemini_client
    if gemini_client is not None:
        return gemini_client

    with gemini_client_lock:
        if gemini_client is not None:
            return gemini_client

        try:
            from django.conf import settings
            api_key = getattr(settings, 'GEMINI_API_KEY', None) or os.environ.get('GEMINI_API_KEY')
        except Exception:
            api_key = os.environ.get('GEMINI_API_KEY')

        if not api_key:
            logger.warning("GEMINI_API_KEY not configured — Gemini extraction disabled, using regex fallback.")
            gemini_client = None
            return None

        try:
            from google.genai import Client
            gemini_client = Client(api_key=api_key)
            return gemini_client
        except Exception as e:
            logger.error(f"Failed to initialise Gemini client: {e}")
            gemini_client = None
            return None


GEMINI_EXTRACTION_PROMPT = (
    "Extract candidate information from this resume text into valid JSON with keys: "
    "name, email, skills (list), experience_years (int), education (list). "
    "Return ONLY the JSON object, no markdown fences, no commentary.\n\n"
    "Resume text:\n{text}"
)


def _clean_json_response(raw_text: str) -> str:
    """Strip markdown code fences if the model wraps its JSON despite instructions."""
    cleaned = raw_text.strip()
    if cleaned.startswith('```'):
        cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned)
        cleaned = re.sub(r'\s*```$', '', cleaned)
    return cleaned.strip()


def extract_with_gemini(text_original: str) -> Optional[Dict]:
    """
    Structured resume-field extraction via the Gemini API.

    Returns a dict with keys: name, email, skills, experience_years, education
    on success, or None on any failure (missing key, SDK error, timeout,
    malformed JSON) — callers must fall back to the regex extractor.
    Never raises.
    """
    client = _get_gemini_client()
    if client is None:
        return None

    if not text_original or not text_original.strip():
        return None

    # Cap input length — long resumes don't need more than ~6000 chars of
    # context for field extraction, and it keeps latency/cost predictable.
    trimmed_text = text_original.strip()[:6000]
    prompt = GEMINI_EXTRACTION_PROMPT.format(text=trimmed_text)

    for model_name in (GEMINI_MODEL_PRIMARY, GEMINI_MODEL_FALLBACK):
        try:
            response = client.models.generate_content(
                model=model_name,
                contents=prompt,
            )

            raw_text = getattr(response, 'text', None)
            if not raw_text:
                logger.warning(f"Gemini ({model_name}) returned an empty response.")
                continue

            parsed = json.loads(_clean_json_response(raw_text))

            # Normalise / validate shape so downstream code can rely on it
            result = {
                'name': parsed.get('name') or None,
                'email': parsed.get('email') or None,
                'skills': parsed.get('skills') if isinstance(parsed.get('skills'), list) else [],
                'experience_years': float(parsed.get('experience_years') or 0),
                'education': parsed.get('education') if isinstance(parsed.get('education'), list) else (
                    [parsed.get('education')] if parsed.get('education') else []
                ),
            }
            logger.info(f"Gemini extraction succeeded using {model_name}.")
            return result

        except json.JSONDecodeError as e:
            logger.error(f"Gemini ({model_name}) returned invalid JSON: {e}")
            continue
        except Exception as e:
            logger.error(f"Gemini extraction failed on {model_name}: {e}")
            continue

    # Both models failed
    return None


# ==============================
# Adaptive Weight System
# ==============================

# Default weights and learning rate
DEFAULT_WEIGHTS = {
    'semantic_weight': 0.6,      # Semantic score (role match) weight
    'analytical_weight': 0.4,    # Analytical score (consistency, skills, learning) weight
}

LEARNING_RATE = 0.05  # eta for weight updates

# Current adaptive weights (loaded/persisted to file)
current_weights = DEFAULT_WEIGHTS.copy()
weights_lock = threading.Lock()  # Thread-safe updates

WEIGHTS_FILE = os.path.join(os.path.dirname(__file__), 'adaptive_weights.json')


def _load_weights_from_disk():
    """Load adaptive weights from disk if available."""
    global current_weights
    if os.path.exists(WEIGHTS_FILE):
        try:
            with open(WEIGHTS_FILE, 'r') as f:
                loaded = json.load(f)
                current_weights = loaded
                logger.info(f"Loaded adaptive weights from disk: {current_weights}")
        except Exception as e:
            logger.error(f"Failed to load weights file: {e}. Using defaults.")
            current_weights = DEFAULT_WEIGHTS.copy()
    else:
        current_weights = DEFAULT_WEIGHTS.copy()


def _save_weights_to_disk():
    """Persist current weights to disk."""
    try:
        with open(WEIGHTS_FILE, 'w') as f:
            json.dump(current_weights, f, indent=2)
        logger.info(f"Saved adaptive weights to disk: {current_weights}")
    except Exception as e:
        logger.error(f"Failed to save weights file: {e}")


def get_current_weights() -> Dict[str, float]:
    """Get the current adaptive weights."""
    with weights_lock:
        return current_weights.copy()


def update_weights(recruiter_decision: str, match_score: float = None):
    """
    Update adaptive weights based on recruiter decision.

    Formula: w = w + (eta * signal)
    - decision in ['HIRE', 'ACCEPT'] -> signal = 1.0 (reinforce)
    - decision in ['REJECT'] -> signal = -1.0 (penalize)
    - decision in ['WAITLIST', 'HOLD'] -> signal = 0.0 (neutral)

    Then normalize so weights sum to 1.0.

    Args:
        recruiter_decision: Decision type ('HIRE', 'ACCEPT', 'REJECT', 'WAITLIST', etc.)
        match_score: Optional match_score to influence which weight gets adjusted more
    """
    global current_weights

    decision_normalized = recruiter_decision.strip().upper()

    # Map decision to signal
    if decision_normalized in ['HIRE', 'ACCEPT']:
        signal = 1.0
    elif decision_normalized in ['REJECT']:
        signal = -1.0
    else:
        # WAITLIST, HOLD, etc. don't update weights
        logger.info(f"Decision '{decision_normalized}' does not trigger weight update.")
        return

    with weights_lock:
        # Adjust weights based on signal
        adjustment = LEARNING_RATE * signal

        # Update both weights with the same adjustment
        current_weights['semantic_weight'] += adjustment
        current_weights['analytical_weight'] += adjustment

        # Clamp to reasonable bounds [0.1, 0.9] to prevent extreme skewing
        current_weights['semantic_weight'] = np.clip(current_weights['semantic_weight'], 0.1, 0.9)
        current_weights['analytical_weight'] = np.clip(current_weights['analytical_weight'], 0.1, 0.9)

        # Normalize so weights sum to 1.0
        total = current_weights['semantic_weight'] + current_weights['analytical_weight']
        current_weights['semantic_weight'] /= total
        current_weights['analytical_weight'] /= total

        # Ensure sum is exactly 1.0 (handle floating point rounding)
        current_weights['analytical_weight'] = round(1.0 - current_weights['semantic_weight'], 6)

        # Log the update
        logger.info(
            f"Weight update from '{decision_normalized}' (signal={signal}): "
            f"semantic={current_weights['semantic_weight']:.4f}, "
            f"analytical={current_weights['analytical_weight']:.4f}"
        )

        # Persist to disk
        _save_weights_to_disk()


# Load weights on module initialization
_load_weights_from_disk()

# Skill Database (Expanded)

SKILL_KEYWORDS = [
    # Programming Languages
    "python", "java", "javascript", "typescript", "c++", "c#", "php", "ruby",
    "go", "golang", "rust", "swift", "kotlin", "scala", "r", "matlab",

    # Web Technologies
    "react", "angular", "vue", "vue.js", "node.js", "express", "django",
    "flask", "fastapi", "spring", "spring boot", "asp.net", "laravel",
    "html", "css", "sass", "less", "bootstrap", "tailwind", "jquery",

    # Databases
    "sql", "mysql", "postgresql", "mongodb", "redis", "elasticsearch",
    "cassandra", "dynamodb", "oracle", "sql server", "sqlite", "mariadb",

    # Cloud & DevOps
    "aws", "azure", "gcp", "google cloud", "docker", "kubernetes", "k8s",
    "jenkins", "gitlab", "github actions", "terraform", "ansible", "ci/cd",

    # Data Science & ML
    "machine learning", "deep learning", "neural networks", "tensorflow",
    "pytorch", "keras", "scikit-learn", "pandas", "numpy", "data analysis",
    "data science", "nlp", "computer vision", "ai", "artificial intelligence",

    # Mobile Development
    "android", "ios", "react native", "flutter", "xamarin", "mobile development",

    # Other Technologies
    "git", "linux", "bash", "powershell", "rest api", "graphql", "microservices",
    "agile", "scrum", "jira", "testing", "unit testing", "selenium", "pytest",
    "api", "backend", "frontend", "full stack", "devops", "blockchain", "solidity"
]

LEARNING_WORDS = [
    "learned", "built", "developed", "improved", "optimized",
    "created", "designed", "implemented", "architected", "led",
    "managed", "delivered", "achieved", "spearheaded", "initiated",
    "enhanced", "streamlined", "automated", "deployed", "scaled"
]


# ==============================
# Sentence Transformer Utilities
# ==============================

def get_similarity_score(resume_text: str, job_description_text: str) -> float:
    """
    Compute semantic similarity between resume and job description using Sentence Transformers.

    Args:
        resume_text: Full resume or job seeker profile text
        job_description_text: Job description or target role text

    Returns:
        Similarity score as a percentage (0-100)
    """
    model = _get_semantic_model()

    if not model:
        logger.warning("Sentence Transformer model not available, returning default score")
        return 50.0

    if not resume_text or not job_description_text:
        return 50.0

    # Cap input length before encoding — mirrors the Gemini extraction cap.
    # A very long resume/JD creates a much bigger tensor during tokenization,
    # which spikes peak memory on exactly the requests most likely to push
    # a memory-constrained instance (e.g. Render's smaller tiers) over its limit.
    resume_text = resume_text.strip()[:6000]
    job_description_text = job_description_text.strip()[:6000]

    try:
        # Encode both texts to embeddings
        resume_embedding = model.encode(resume_text, convert_to_numpy=True)
        job_embedding = model.encode(job_description_text, convert_to_numpy=True)

        # Compute cosine similarity and convert to percentage
        similarity = cosine_similarity([resume_embedding], [job_embedding])[0][0]
        return float(similarity * 100)
    except Exception as e:
        logger.error(f"Error computing similarity score: {e}")
        return 50.0


# Email and phone regex patterns
EMAIL_PATTERN = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
PHONE_PATTERN = re.compile(r'[\+]?[(]?[0-9]{1,4}[)]?[-\s\.]?[(]?[0-9]{1,4}[)]?[-\s\.]?[0-9]{1,4}[-\s\.]?[0-9]{1,9}')

# Education keywords with levels
EDUCATION_KEYWORDS = {
    "phd": "PhD",
    "doctorate": "PhD",
    "ph.d": "PhD",
    "master": "Master's Degree",
    "m.tech": "Master's Degree",
    "m.s": "Master's Degree",
    "msc": "Master's Degree",
    "mba": "MBA",
    "bachelor": "Bachelor's Degree",
    "b.tech": "Bachelor's Degree",
    "b.e": "Bachelor's Degree",
    "b.s": "Bachelor's Degree",
    "bsc": "Bachelor's Degree",
    "associate": "Associate Degree",
    "diploma": "Diploma"
}


# ==============================
# Text Extraction
# ==============================

def extract_text_from_pdf(file_path: str) -> Tuple[str, str]:
    """
    Extract text from PDF or DOCX file.
    Returns: (lowercase_text, original_case_text)
    Raises: Exception if file cannot be read
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = os.path.splitext(file_path.lower())[1]

    # ── DOCX handling ──────────────────────────────────────────────
    if ext in ('.docx', '.doc'):
        try:
            import docx as python_docx
            doc = python_docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            text = "\n".join(paragraphs)
            if not text.strip():
                raise ValueError("DOCX contains no extractable text")
            return text.lower(), text
        except ImportError:
            raise ValueError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {str(e)}")
            raise

    # ── PDF handling ───────────────────────────────────────────────
    if ext == '.pdf':
        try:
            import fitz as pymupdf
            text = ""
            doc = pymupdf.open(file_path)
            for page in doc:
                text += page.get_text()
            doc.close()
            if not text.strip():
                raise ValueError("PDF contains no extractable text (may be scanned image)")
            return text.lower(), text
        except ImportError:
            raise ValueError("PyMuPDF not installed. Run: pip install PyMuPDF")
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {str(e)}")
            raise

    raise ValueError(f"Unsupported file type: {ext}. Please upload PDF or DOCX.")


# ==============================
# Core Extractors
# ==============================

def extract_skills(text: str) -> List[str]:
    """Extract skills from resume text with deduplication."""
    found_skills = set()
    text_lower = text.lower()

    for skill in SKILL_KEYWORDS:
        # Use word boundaries for better matching
        if re.search(r'\b' + re.escape(skill) + r'\b', text_lower):
            found_skills.add(skill)

    return sorted(list(found_skills))


def extract_email(text: str) -> Optional[str]:
    """Extract email address from resume."""
    emails = EMAIL_PATTERN.findall(text)
    # Return first valid email found
    return emails[0] if emails else None


def extract_phone(text: str) -> Optional[str]:
    """Extract phone number from resume."""
    phones = PHONE_PATTERN.findall(text)
    # Return first phone that looks valid (at least 10 digits)
    for phone in phones:
        digits = re.sub(r'\D', '', phone)
        if len(digits) >= 10:
            return phone
    return None


def extract_name(text: str) -> Optional[str]:
    """
    Improved name extraction using multiple strategies.
    Looks at first few lines and validates against common patterns.
    """
    lines = text.split("\n")

    # Common resume headers to skip
    skip_keywords = [
        'resume', 'cv', 'curriculum vitae', 'profile', 'objective',
        'contact', 'email', 'phone', 'address', 'linkedin'
    ]

    for line in lines[:15]:  # Check first 15 lines
        clean = line.strip()
        clean_lower = clean.lower()

        # Skip empty lines and headers
        if not clean or any(skip in clean_lower for skip in skip_keywords):
            continue

        # Check if line looks like a name
        words = clean.split()
        if 2 <= len(words) <= 4:  # Names are usually 2-4 words
            # No digits, reasonable length, mostly alphabetic
            if (not any(char.isdigit() for char in clean) and
                    len(clean) < 40 and
                    sum(c.isalpha() or c.isspace() for c in clean) / len(clean) > 0.8):
                return clean.title()

    return None


def extract_experience(text: str) -> float:
    """
    Extract years of experience from resume.
    Looks for various patterns like "5 years", "5+ years", etc.
    """
    # Pattern for explicit years mention
    patterns = [
        r'(\d+)\+?\s*(?:years?|yrs?)\s+(?:of\s+)?experience',
        r'experience\s*:\s*(\d+)\+?\s*(?:years?|yrs?)',
        r'(\d+)\+?\s*(?:years?|yrs?)',
    ]

    max_years = 0
    for pattern in patterns:
        matches = re.findall(pattern, text.lower())
        if matches:
            max_years = max(max_years, max(int(m) for m in matches))

    return float(max_years) if max_years > 0 else 0.0


def extract_education(text: str) -> Optional[str]:
    """
    Extract highest education level from resume.
    Returns the highest degree found with confidence.
    """
    text_lower = text.lower()
    highest_level = None
    priority = 0  # Higher priority = higher education level

    education_priority = {
        "PhD": 4,
        "Master's Degree": 3,
        "MBA": 3,
        "Bachelor's Degree": 2,
        "Associate Degree": 1,
        "Diploma": 1
    }

    for keyword, degree in EDUCATION_KEYWORDS.items():
        if keyword in text_lower:
            current_priority = education_priority.get(degree, 0)
            if current_priority > priority:
                priority = current_priority
                highest_level = degree

    return highest_level


# ==============================
# Intelligence Scores
# ==============================

def compute_learning_velocity(text: str) -> str:
    """
    Learning velocity label — High / Medium / Low.
    Based on growth-indicator keyword frequency in resume text.
    """
    count = sum(1 for word in LEARNING_WORDS if word in text.lower())
    if count >= 8:
        return "High"
    elif count >= 4:
        return "Medium"
    return "Low"


def compute_learning_velocity_score(skills: List[str], experience_years: float) -> float:
    """
    Learning Velocity numeric score (paper §IV.E3).
    Formula: LV = new_skills_count / max(1, years)
    Normalised 0-100 using typical range cap of 10 skills/year as max.
    """
    if experience_years <= 0:
        experience_years = 1.0
    raw_lv = len(skills) / experience_years
    # Normalise: cap at 10 skills/year as practical maximum
    LV_MAX = 10.0
    LV_MIN = 0.0
    normalised = (raw_lv - LV_MIN) / (LV_MAX - LV_MIN)
    return round(min(100.0, max(0.0, normalised * 100)), 2)


def compute_role_match(resume_text: str, target_role: str, skills: List[str],
                       job_description: str = "") -> float:
    """
    Role-Specific Suitability and Matching Analysis (paper §IV.E).
    Uses Sentence Transformer embeddings + skill overlap boost.
    """
    match_target = job_description.strip() if job_description and job_description.strip() else target_role

    if not match_target:
        return 50.0

    semantic_score = get_similarity_score(resume_text, match_target.lower())

    jd_words = set(match_target.lower().split())
    resume_skills = set(skills)
    skill_overlap = sum(1 for s in resume_skills if any(w in s for w in jd_words))
    skill_bonus = min(30, skill_overlap * 8)

    if job_description and job_description.strip():
        final_score = (0.70 * semantic_score) + (0.30 * (50 + skill_bonus))
    else:
        final_score = (0.60 * semantic_score) + (0.40 * (50 + skill_bonus))

    return round(min(100, max(0, final_score)), 2)


def compute_consistency_score(skills: List[str], experience_years: float) -> float:
    """
    Cross-document consistency analysis (paper §IV.A).
    """
    score = 100.0
    skills_lower = [s.lower() for s in skills]

    senior_skills = [
        'kubernetes', 'terraform', 'system design', 'distributed systems',
        'kafka', 'grpc', 'microservices', 'architect'
    ]
    has_senior = any(s in skills_lower for s in senior_skills)
    if has_senior and experience_years < 1.5:
        score -= 20

    expected_max_skills = max(6, experience_years * 4)
    if len(skills) > expected_max_skills:
        score -= 12

    advanced = {'kubernetes', 'terraform', 'kafka', 'pytorch', 'tensorflow'}
    foundational = {'python', 'javascript', 'java', 'git', 'sql', 'linux'}
    if bool(advanced & set(skills_lower)) and not bool(foundational & set(skills_lower)):
        score -= 15

    if experience_years >= 3 and len(skills) < 4:
        score -= 18

    return round(max(30.0, score), 2)


def validate_skills_from_text(skills: List[str], text_lower: str) -> List[dict]:
    """
    Evidence-backed skill validation (paper §IV.B).
    """
    ACTION_VERBS = [
        'built', 'developed', 'designed', 'implemented', 'deployed',
        'created', 'architected', 'led', 'managed', 'used', 'integrated',
        'automated', 'optimised', 'optimized', 'migrated', 'scaled',
        'wrote', 'delivered', 'maintained', 'configured', 'set up',
    ]
    validated = []
    for skill in skills:
        sl = skill.lower()
        idx = text_lower.find(sl)
        if idx == -1:
            validated.append({'skill': skill, 'status': 'Unverified'})
            continue
        window_start = max(0, idx - 60)
        window_end = min(len(text_lower), idx + len(sl) + 60)
        window = text_lower[window_start:window_end]
        has_action = any(verb in window for verb in ACTION_VERBS)
        if has_action:
            validated.append({'skill': skill, 'status': 'Valid'})
        else:
            count = text_lower.count(sl)
            if count >= 2:
                validated.append({'skill': skill, 'status': 'Partial'})
            else:
                validated.append({'skill': skill, 'status': 'Unverified'})
    return validated


def compute_skill_validation_score(skills: List[str], learning_velocity: str) -> float:
    """Validate skills based on quantity and learning indicators."""
    base_score = 40
    skill_bonus = min(40, len(skills) * 4)
    velocity_bonus = {"High": 20, "Medium": 10, "Low": 5}.get(learning_velocity, 5)
    total = base_score + skill_bonus + velocity_bonus
    return round(min(100, total), 2)


def compute_resume_strength_score(skills: List[str], education: Optional[str],
                                  experience_years: float) -> float:
    """Overall resume strength based on completeness and quality."""
    base_score = 30
    skill_bonus = min(35, len(skills) * 3)

    education_bonus = 0
    if education:
        education_levels = {
            "PhD": 20,
            "Master's Degree": 15,
            "MBA": 15,
            "Bachelor's Degree": 12,
            "Associate Degree": 8,
            "Diploma": 5
        }
        education_bonus = education_levels.get(education, 5)

    experience_bonus = min(15, experience_years * 2)

    total = base_score + skill_bonus + education_bonus + experience_bonus
    return round(min(100, total), 2)


def compute_analytical_score(consistency_score: float,
                             skill_validation_score: float,
                             lv_score: float) -> float:
    """
    Analytical Score (paper §IV.F): AS = (CS + VS + LV) / 3
    """
    return round((consistency_score + skill_validation_score + lv_score) / 3, 2)


def compute_final_score(semantic_score: float, analytical_score: float,
                        alpha: float = None, beta: float = None) -> float:
    """
    Final Score (paper §IV.G) with adaptive weights: FS = alpha * SS + beta * AS
    """
    if alpha is None or beta is None:
        weights = get_current_weights()
        alpha = weights['semantic_weight']
        beta = weights['analytical_weight']

    return round(alpha * semantic_score + beta * analytical_score, 2)


def compute_final_fit(final_score: float) -> str:
    """
    Classify final fit from the combined Final Score.
    Thresholds aligned to paper: High >= 70, Medium >= 40, Low < 40.
    """
    if final_score >= 70:
        return "High"
    elif final_score >= 40:
        return "Medium"
    return "Low"


def generate_career_trajectory(skills: List[str], learning_velocity: str,
                               experience_years: float) -> str:
    """Generate a meaningful career trajectory summary."""
    skill_count = len(skills)

    if skill_count >= 15:
        skill_level = "extensive"
    elif skill_count >= 8:
        skill_level = "strong"
    else:
        skill_level = "developing"

    return (
        f"Candidate demonstrates {skill_level} technical expertise with "
        f"{skill_count} identified skills, {learning_velocity.lower()} learning velocity, "
        f"and {experience_years:.1f} years of experience."
    )


# ==============================
# ✅ MAIN ANALYSIS FUNCTION
# ==============================

def analyze_resume(file_path: str, target_role: str = "", job_description: str = "") -> Dict:
    """
    Main function to analyze a resume PDF/DOCX and extract comprehensive insights.

    Extraction strategy: try Gemini structured extraction first (name, email,
    skills, experience_years, education). If it fails for any reason — no API
    key, SDK error, timeout, malformed JSON — silently fall back to the
    original regex-based extractors. Phone and all scoring always run
    against the locally extracted text regardless of which path was used,
    since Gemini's schema doesn't include phone and the scoring functions
    need the raw text either way.

    Args:
        file_path: Path to the resume file (PDF or DOCX)
        target_role: Target job role for matching (optional)
        job_description: Full job description text for richer matching (optional)

    Returns:
        Dict containing all analysis results

    Raises:
        FileNotFoundError: If file doesn't exist
        ValueError: If file is not a PDF/DOCX or contains no text
        Exception: For other processing errors
    """
    try:
        # Extract text (both lowercase and original case) — always needed,
        # regardless of extraction path, for scoring and phone lookup.
        text_lower, text_original = extract_text_from_pdf(file_path)

        # ── Try Gemini structured extraction first ──────────────────────
        gemini_result = extract_with_gemini(text_original)

        if gemini_result:
            name = gemini_result['name']
            email = gemini_result['email']
            skills = [s.lower() for s in gemini_result['skills']] or extract_skills(text_lower)
            experience_years = gemini_result['experience_years'] or extract_experience(text_lower)
            education_list = gemini_result['education']
            education = education_list[0] if education_list else extract_education(text_lower)
            extraction_method = 'gemini'
        else:
            name = extract_name(text_original)
            email = extract_email(text_original)
            skills = extract_skills(text_lower)
            experience_years = extract_experience(text_lower)
            education = extract_education(text_lower)
            extraction_method = 'regex'

        # Phone isn't part of the Gemini schema — always regex
        phone = extract_phone(text_original)

        # Learning velocity label still runs off raw text keyword frequency
        learning_velocity = compute_learning_velocity(text_lower)

        # ── Semantic Similarity Score (SS) — paper §IV.C/D ──────────────
        role_score = compute_role_match(text_lower, target_role, skills, job_description)

        # ── Analytical Score components — paper §IV.E ───────────────────
        consistency_score = compute_consistency_score(skills, experience_years)
        resume_strength_score = compute_resume_strength_score(skills, education, experience_years)

        validated_skills = validate_skills_from_text(skills, text_lower)
        valid_count = sum(1 for v in validated_skills if v['status'] == 'Valid')
        skill_validation_score = (
            round((valid_count / len(validated_skills)) * 100, 2)
            if validated_skills else
            compute_skill_validation_score(skills, learning_velocity)
        )

        lv_score = compute_learning_velocity_score(skills, experience_years)

        analytical_score = compute_analytical_score(
            consistency_score, skill_validation_score, lv_score
        )

        final_score = compute_final_score(role_score, analytical_score)
        final_fit = compute_final_fit(final_score)

        career_trajectory = generate_career_trajectory(skills, learning_velocity, experience_years)

        missing_skills = [s for s in SKILL_KEYWORDS[:50] if s not in skills][:8]

        jd_note = " Matched against full job description." if job_description and job_description.strip() else ""
        explainability = (
            f"Semantic match: {role_score:.0f}%.{jd_note} "
            f"Analytical score: {analytical_score:.0f}/100 "
            f"(Consistency: {consistency_score:.0f}, Skill validation: {skill_validation_score:.0f}, "
            f"Learning velocity: {lv_score:.0f}). "
            f"Final score: {final_score:.0f}/100. "
            f"Profile shows {learning_velocity.lower()} learning velocity with "
            f"{experience_years:.1f} years of experience."
        )

        profile_summary = (
            f"{name or 'Candidate'} - {education or 'Education not specified'} | "
            f"{len(skills)} skills | {experience_years:.1f} years experience"
        )

        logger.info(f"Successfully analyzed resume: {name or 'Unknown'} (extraction: {extraction_method})")

        return {
            # Basic Information
            "name": name,
            "email": email,
            "phone": phone,

            # Skills and Learning
            "skills": skills,
            "learning_velocity": learning_velocity,
            "missing_skills": missing_skills,

            # Experience and Education
            "experience_years": experience_years,
            "education": education,

            # Semantic similarity (paper §IV.C)
            "role_match_score": role_score,

            # Analytical score components (paper §IV.E-F)
            "consistency_score": consistency_score,
            "skill_validation_score": skill_validation_score,
            "lv_score": lv_score,
            "analytical_score": analytical_score,

            # Final combined score (paper §IV.G: FS = 0.5*SS + 0.5*AS)
            "final_score": final_score,
            "final_fit": final_fit,

            # Supplementary
            "resume_strength_score": resume_strength_score,

            # Insights
            "career_trajectory": career_trajectory,
            "profile_summary": profile_summary,
            "explainability": explainability,

            # Evidence-backed skill validation (paper §IV.B)
            "validated_skills": validated_skills,

            # Which extraction path produced name/email/skills/experience/education
            "extraction_method": extraction_method,
        }

    except FileNotFoundError:
        logger.error(f"File not found: {file_path}")
        raise
    except ValueError as e:
        logger.error(f"Invalid file or content: {str(e)}")
        raise
    except Exception as e:
        logger.error(f"Unexpected error during resume analysis: {str(e)}")
        raise Exception(f"Resume analysis failed: {str(e)}")


# ==============================
# GitHub Candidate Search Service
# ==============================

GITHUB_API_BASE = "https://api.github.com"
GITHUB_SEARCH_USERS = f"{GITHUB_API_BASE}/search/users"
GITHUB_USER_REPOS = f"{GITHUB_API_BASE}/users/{{username}}/repos"
GITHUB_USER_PROFILE = f"{GITHUB_API_BASE}/users/{{username}}"

GITHUB_REQUEST_TIMEOUT = 10
GITHUB_RATE_LIMIT_DELAY = 1
MAX_REPOS_TO_ANALYZE = 30


def _make_github_request(url: str, params: Optional[Dict] = None,
                         timeout: int = GITHUB_REQUEST_TIMEOUT) -> Optional[Dict]:
    try:
        headers = {
            'Accept': 'application/vnd.github.v3+json',
            'User-Agent': 'NeuroHire-Recruiter-Platform'
        }

        github_token = os.environ.get('GITHUB_TOKEN')
        if github_token:
            headers['Authorization'] = f'token {github_token}'

        response = requests.get(
            url,
            params=params,
            headers=headers,
            timeout=timeout
        )

        if response.status_code == 403:
            rate_limit_remaining = response.headers.get('X-RateLimit-Remaining', '0')
            if rate_limit_remaining == '0':
                reset_time = response.headers.get('X-RateLimit-Reset', 'unknown')
                logger.warning(f"GitHub API rate limit exceeded. Resets at: {reset_time}")
                return None

        response.raise_for_status()
        time.sleep(GITHUB_RATE_LIMIT_DELAY)
        return response.json()

    except requests.exceptions.Timeout:
        logger.error(f"GitHub API request timeout for URL: {url}")
        return None
    except requests.exceptions.RequestException as e:
        logger.error(f"GitHub API request failed: {str(e)}")
        return None
    except ValueError as e:
        logger.error(f"Failed to parse GitHub API response: {str(e)}")
        return None


def _extract_languages_from_repos(username: str, max_repos: int = MAX_REPOS_TO_ANALYZE) -> Dict[str, int]:
    languages_aggregate = {}

    try:
        repos_url = GITHUB_USER_REPOS.format(username=username)
        params = {
            'sort': 'updated',
            'direction': 'desc',
            'per_page': max_repos
        }

        repos_data = _make_github_request(repos_url, params)
        if not repos_data:
            return languages_aggregate

        for repo in repos_data[:max_repos]:
            if repo.get('fork', False):
                continue

            language = repo.get('language')
            if language:
                languages_aggregate[language] = languages_aggregate.get(language, 0) + 1

        return languages_aggregate

    except Exception as e:
        logger.error(f"Failed to extract languages for user {username}: {str(e)}")
        return languages_aggregate


def _get_top_languages(languages_dict: Dict[str, int], top_n: int = 5) -> List[str]:
    if not languages_dict:
        return []

    sorted_languages = sorted(
        languages_dict.items(),
        key=lambda x: x[1],
        reverse=True
    )

    return [lang for lang, _ in sorted_languages[:top_n]]


def _format_github_user(user_data: Dict, languages: List[str]) -> Dict:
    return {
        'username': user_data.get('login', ''),
        'profile_url': user_data.get('html_url', ''),
        'avatar_url': user_data.get('avatar_url', ''),
        'bio': user_data.get('bio', '') or 'No bio available',
        'name': user_data.get('name', '') or user_data.get('login', ''),
        'location': user_data.get('location', ''),
        'email': user_data.get('email', ''),
        'company': user_data.get('company', ''),
        'public_repos': user_data.get('public_repos', 0),
        'followers': user_data.get('followers', 0),
        'following': user_data.get('following', 0),
        'created_at': user_data.get('created_at', ''),
        'top_languages': languages,
        'profile_score': _calculate_github_profile_score(user_data, languages),
    }


def _calculate_github_profile_score(user_data: Dict, languages: List[str]) -> float:
    score = 0.0

    repos = user_data.get('public_repos', 0)
    score += min(30, repos * 1.5)

    followers = user_data.get('followers', 0)
    score += min(25, followers * 0.5)

    score += min(20, len(languages) * 4)

    if user_data.get('bio'):
        score += 10

    if user_data.get('company'):
        score += 8

    if user_data.get('location'):
        score += 7

    return round(min(100, score), 2)


def github_search_service(query: str, max_results: int = 10) -> Dict:
    try:
        if not query or not query.strip():
            return {
                'success': False,
                'error': 'Search query is required',
                'data': [],
                'total_count': 0,
                'query': query
            }

        max_results = min(max_results, 100)

        logger.info(f"GitHub search initiated: query='{query}', max_results={max_results}")

        search_params = {
            'q': query,
            'per_page': max_results,
            'sort': 'followers',
            'order': 'desc'
        }

        search_response = _make_github_request(GITHUB_SEARCH_USERS, search_params)

        if not search_response:
            return {
                'success': False,
                'error': 'GitHub API request failed or rate limit exceeded',
                'data': [],
                'total_count': 0,
                'query': query
            }

        total_count = search_response.get('total_count', 0)
        users = search_response.get('items', [])

        logger.info(f"Found {total_count} users, processing {len(users)} results")

        formatted_users = []
        for user in users:
            username = user.get('login')
            if not username:
                continue

            profile_url = GITHUB_USER_PROFILE.format(username=username)
            user_details = _make_github_request(profile_url)

            if not user_details:
                user_details = user

            languages_dict = _extract_languages_from_repos(username)
            top_languages = _get_top_languages(languages_dict)

            formatted_user = _format_github_user(user_details, top_languages)
            formatted_users.append(formatted_user)

            logger.info(f"Processed user: {username} ({len(top_languages)} languages)")

        logger.info(f"GitHub search completed successfully: {len(formatted_users)} users processed")

        return {
            'success': True,
            'data': formatted_users,
            'total_count': total_count,
            'query': query,
            'results_returned': len(formatted_users)
        }

    except Exception as e:
        logger.error(f"GitHub search service failed: {str(e)}")
        return {
            'success': False,
            'error': f'Search failed: {str(e)}',
            'data': [],
            'total_count': 0,
            'query': query
        }